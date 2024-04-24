import html
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

from django.template.defaultfilters import date as format_date
from django.utils.text import capfirst, normalize_newlines

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Mm, Pt

from .constants import ExportModes, Leerjaren
from .models import Toets, Vak

LEERJAAR_WEGING = {
    Leerjaren.gym_ath_3: ("Weging", "weging_r4"),
    Leerjaren.havo_4: ("Weging SE", "weging_ed4"),
    Leerjaren.vwo_4: ("Weging SE", "weging_ed4"),
    Leerjaren.havo_5: ("Weging SE", "weging_ed5"),
    Leerjaren.vwo_5: ("Weging SE", "weging_ed5"),
    Leerjaren.vwo_6: ("Weging SE", "weging_ed6"),
}

HIDE_DOMEIN = {Leerjaren.gym_ath_3}


R4_LEERJAREN = (
    Leerjaren.vwo_4,
    Leerjaren.vwo_5,
    Leerjaren.havo_4,
)

HEADER_BG_COLOR = "D9D9D9"

COLUMN_WIDTHS = {
    0: Cm(1.51),
    1: Cm(10.43),
    2: Cm(3.25),
    3: Cm(1.75),
    4: Cm(1.5),
    5: Cm(2.25),
    6: Cm(1.5),
}

COLUMN_WIDTHS_PER_LEERJAAR = {
    Leerjaren.gym_ath_3: {
        0: Cm(1.51),
        1: Cm(10.43),
        2: Cm(1.75),
        3: Cm(1.5),
        4: Cm(2.75),
        5: Cm(1.5),
        6: Cm(2.0),
    }
}


def initialize_document() -> Document:
    document = Document()

    # set to landscape
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    # A4 page size
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    # set correct margins
    section.bottom_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    return document


def add_vak(
    document: Document,
    logo_path: str,
    vak: Vak,
    year: int,
    leerjaar: int,
    toetsweek_periodes: dict[int, list[int]],
) -> None:
    match leerjaar:
        # Klas 6
        case Leerjaren.overstappers_vwo_5:
            add_vak_overstappers_vwo5(document, logo_path, vak, year, leerjaar)
        # Klas 7
        case Leerjaren.overstappers_vwo_6:
            add_vak_overstappers_vwo6(document, logo_path, vak, year, leerjaar)
        case Leerjaren.tl_3 | Leerjaren.tl_4:
            add_vak_tl(document, vak, year, leerjaar, toetsweek_periodes)
        # Default
        case _:
            add_vak_regular(
                document, logo_path, vak, year, leerjaar, toetsweek_periodes
            )


@dataclass
class Omschrijving:
    content: str
    inleverdatum: str
    voetnoot: str

    _voetnoot_counter = 0
    voetnoot_nr = 0

    def __post_init__(self):
        if self.voetnoot:
            self.__class__._voetnoot_counter += 1
            self.voetnoot_nr = self.__class__._voetnoot_counter

    @classmethod
    def reset_counter(cls):
        cls._voetnoot_counter = 0


def clean_text(text: str) -> str:
    return html.unescape(normalize_newlines(text).strip())


def get_periode_and_week(
    toets: Toets, toetsweek_periodes: dict[int, list[int]]
) -> tuple[str, str]:
    toetsweken = sum(toetsweek_periodes.values(), [])

    periode = toets.periode
    week = toets.week or ""

    if toets.week in toetsweken:
        periode = f"{periode} (tw)"
        week = "/".join([str(wk) for wk in toetsweek_periodes[toets.periode]])

    return str(periode), week


def get_toets_table(
    leerjaar: int,
    vak: Vak,
    toetsweek_periodes: dict[int, list[int]],
    weging: Optional[Tuple[str, str]],
) -> List[List[str]]:
    header = [
        "Code",
        "Onderwerp/Omschrijving",
        "Domein",
        "Periode",
        "Week",
        "Soort werk",
        "Tijd\n(min)",
    ]

    if leerjaar in HIDE_DOMEIN:
        header.remove("Domein")

    if leerjaar in R4_LEERJAREN:
        header.append("Weging R4")

    if weging is not None:
        header.append(weging[0])

    rows = []

    for toets in vak.toetsen:
        periode, week = get_periode_and_week(toets, toetsweek_periodes)

        if toets.inleverdatum:
            formatted = format_date(toets.inleverdatum, "j F Y")
            inleverdatum = f"inleverdatum: {formatted}"
        elif toets.datum:
            formatted = format_date(toets.datum, "j F Y")
            inleverdatum = f"datum: {formatted}"
        else:
            inleverdatum = ""

        soort_werk = (
            toets.soortwerk.naam if toets.soortwerk and toets.soortwerk.id != 12 else ""
        )
        row = [
            toets.code,
            Omschrijving(
                content=clean_text(toets.omschrijving),
                inleverdatum=inleverdatum,
                voetnoot=toets.voetnoot,
            ),
            periode or "",
            week,
            soort_werk,
            toets.tijd or "",
        ]

        if leerjaar not in HIDE_DOMEIN:
            row.insert(2, toets.domein or "")

        if leerjaar in R4_LEERJAREN:
            row.append(toets.weging_r4 or "")

        if weging is not None:
            row.append(getattr(toets, weging[1]) or "")

        assert len(row) == len(header), "Header and row columns mismatch"
        rows.append(row)

    return [header] + rows


def add_header(document: Document, vak: Vak, year: int, leerjaar: int):
    _leerjaar = Leerjaren(leerjaar).label
    school_year = f"{year}-{year + 1}"

    # set up the table header
    header_p = document.add_paragraph()
    header_tab_stops = header_p.paragraph_format.tab_stops
    header_tab_stops.add_tab_stop(Cm(3))
    header_tab_stops.add_tab_stop(Cm(22.71), alignment=WD_TAB_ALIGNMENT.RIGHT)
    header_tab_stops.add_tab_stop(Cm(23.06))

    vak_naam = capfirst(html.unescape(vak.naam))
    if vak.afkorting and vak.weergeven == 1:
        vak_naam = f"{vak_naam} ({vak.afkorting})"

    prefix = "PTB" if leerjaar == Leerjaren.gym_ath_3 else "PTA"
    header_run = header_p.add_run(f"{prefix}\t{vak_naam}\t{_leerjaar}\t{school_year}")
    # header_run.font.name = "Arial"
    header_run.font.size = Pt(14)
    header_run.bold = True


def add_vak_regular(
    document: Document,
    logo_path: str,
    vak: Vak,
    year: int,
    leerjaar: int,
    toetsweek_periodes: dict[int, List[int]],
):
    if not vak.toetsen:
        return

    section = document.sections[-1]

    weging = LEERJAAR_WEGING.get(leerjaar)

    add_header(document, vak, year, leerjaar)

    # add the logo
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # paragraph.add_run().add_picture(logo_path, width=Cm(5.68))

    # try to set global font name
    _set_default_font(paragraph)

    Omschrijving.reset_counter()
    [header, *rows] = get_toets_table(leerjaar, vak, toetsweek_periodes, weging)

    table = document.add_table(rows=len(rows) + 1, cols=len(header))
    _set_default_table_style(table)

    header_cells = table.rows[0].cells
    for index, title in enumerate(header):
        cell = header_cells[index]
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell)

    toets_content = []

    for index, row in enumerate(rows, 1):
        row_cells = table.rows[index].cells
        for _index, content in enumerate(row):
            cell = row_cells[_index]
            par = cell.paragraphs[0]

            # normal content
            if not isinstance(content, Omschrijving):
                par.text = str(content)
            # calculated content
            else:
                toets_content.append(content)
                par.add_run(text=content.content)

                if content.voetnoot:
                    par.add_run(text=" ")
                    voetnoot_run = par.add_run(text=f"{content.voetnoot_nr})")
                    voetnoot_run.font.superscript = True

                if content.inleverdatum:
                    datum_par = cell.add_paragraph(text=content.inleverdatum)
                    datum_par.runs[0].font.size = Pt(8)
                    datum_par.runs[0].italic = True
    _style_table_cells(table)

    column_widths = COLUMN_WIDTHS_PER_LEERJAAR.get(leerjaar, COLUMN_WIDTHS)

    # set the column widths
    num_extra_columns = len(table.columns) - len(column_widths)
    remaining_width = (
        section.page_width
        - section.left_margin
        - section.right_margin
        - sum(column_widths.values())
        - Mm(1)
    )
    extra_column_width = (
        int(remaining_width / num_extra_columns) if num_extra_columns else 0
    )
    for index, column in enumerate(table.columns):
        column.width = column_widths.get(index, extra_column_width)
        # sigh... dumb format
        for cell in column.cells:
            cell.width = column.width

    # REMARKS, below the table
    toets_voetnoot_content = [x for x in toets_content if x.voetnoot]
    vak_voetnoten = [voetnoot.noot for voetnoot in vak.voetnoten]

    if toets_voetnoot_content or vak_voetnoten:
        remarks_p = document.add_paragraph("opmerking")
        remarks_p.runs[0].bold = True
        remarks_p.paragraph_format.space_before = Pt(10)

    if toets_voetnoot_content:
        p_toets_voetnoten = document.add_paragraph()
        p_toets_voetnoten.paragraph_format.space_after = 0
        for content in toets_voetnoot_content:
            run_super = p_toets_voetnoten.add_run(text=f"{content.voetnoot_nr})")
            run_super.font.superscript = True
            p_toets_voetnoten.add_run(text=f" {normalize_newlines(content.voetnoot)}\n")

    if vak_voetnoten:
        voetnoten_p = document.add_paragraph()
        voetnoten_p.paragraph_format.space_after = 0
        for voetnoot in vak_voetnoten:
            _voetnoot = normalize_newlines(voetnoot)
            voetnoten_p.add_run(f"{_voetnoot}\n")

    document.add_page_break()


def iter_oude_toets_columns(oude_toets: Optional[Toets]):
    if oude_toets is None:
        yield "-"
        yield "-"
        yield "(toets ontbreekt)"
        yield "-"
        return

    yield oude_toets.code,
    yield f"{oude_toets.jaar}-{oude_toets.jaar + 1}",
    yield clean_text(oude_toets.omschrijving),
    yield oude_toets.domein,


def add_vak_overstappers_vwo5(
    document: Document,
    logo_path: str,
    vak: Vak,
    year: int,
    leerjaar: int,
):
    match vak.export_bit:
        case ExportModes.no_export.value:
            return
        case ExportModes.table.value:
            render_vak = any(
                (
                    vak.overnemen_herwaarderen,  # type: ignore
                    vak.inhalen,  # type: ignore
                    vak.inhaalopdrachten,  # type: ignore
                )
            )
            if not render_vak:
                return
            # otherwise, do the old behaviour after the pattern match
        case ExportModes.remark_completed_earlier.value:
            add_header(document, vak, year, leerjaar)
            paragraph = document.add_paragraph(
                "De leerling heeft dit vak in een eerder stadium al afgerond op Vwo "
                "niveau. De eindbeoordeling van dit vak wordt meegenomen."
            )
            # try to set global font name
            _set_default_font(paragraph)
            document.add_page_break()
            return

    add_header(document, vak, year, leerjaar)

    if vak.overnemen_herwaarderen:  # type: ignore

        paragraph = document.add_paragraph(
            "De volgende al gemaakte onderdelen worden meegenomen uit "
            "eerdere leerjaren.\nDaarbij wordt het cijfer overgenomen, of de "
            "toets wordt opnieuw gewaardeerd."
        )
        # try to set global font name
        _set_default_font(paragraph)

        header = [
            "Code V4/V5",
            "Jaar",
            "Omschrijving",
            "Domein",
            "Overnemen/Herwaarderen",
            "Weging SE",
        ]
        table_data = [
            [
                *iter_oude_toets_columns(overstap.oude_toets),
                overstap.get_actie_display(),
                str(overstap.weging_ed4),
            ]
            for overstap in vak.overnemen_herwaarderen
        ]
        WIDTHS = {
            0: Cm(2.00),
            1: Cm(2.50),
            2: Cm(10.43),
            3: Cm(3.25),
            4: Cm(3.00),
            5: Cm(2.00),
        }
        create_table(document, header, table_data, index_no_center=2, widths=WIDTHS)

    if vak.inhalen:
        paragraph = document.add_paragraph(
            "De volgende toetsen uit havo 4 moeten worden ingehaald."
        )
        # try to set global font name
        _set_default_font(paragraph)
        if vak.overnemen_herwaarderen:
            paragraph.paragraph_format.space_before = Pt(10)

        header = ["Code H4", "Jaar", "Omschrijving", "Domein", "Weging SE"]
        table_data = [
            [
                overstap.oude_toets.code,
                f"{overstap.oude_toets.jaar}-{overstap.oude_toets.jaar + 1}",
                clean_text(overstap.oude_toets.omschrijving),
                overstap.oude_toets.domein,
                str(overstap.weging_ed4),
            ]
            for overstap in vak.inhalen
        ]
        WIDTHS = {
            0: Cm(2.00),
            1: Cm(2.50),
            2: Cm(10.43),
            3: Cm(3.25),
            4: Cm(2.00),
        }
        create_table(document, header, table_data, index_no_center=2, widths=WIDTHS)

    if vak.inhaalopdrachten:
        paragraph = document.add_paragraph(
            "Tevens moet de volgende inhaalopdracht worden gemaakt om aan het "
            "schoolexamen Havo te voldoen:"
        )
        # try to set global font name
        _set_default_font(paragraph)
        if vak.overnemen_herwaarderen or vak.inhalen:
            paragraph.paragraph_format.space_before = Pt(10)

        header = ["Code H4", "Jaar", "Omschrijving", "Domein", "Weging SE"]
        table_data = [
            [
                toets.code,
                f"{toets.jaar}-{toets.jaar + 1}",
                clean_text(toets.omschrijving),
                toets.domein,
                str(toets.weging_ed4),
            ]
            for toets in vak.inhaalopdrachten
        ]
        WIDTHS = {
            0: Cm(2.00),
            1: Cm(2.50),
            2: Cm(10.43),
            3: Cm(3.25),
            4: Cm(2.00),
        }
        create_table(document, header, table_data, index_no_center=2, widths=WIDTHS)

    document.add_page_break()


def add_vak_overstappers_vwo6(
    document: Document,
    logo_path: str,
    vak: Vak,
    year: int,
    leerjaar: int,
) -> None:
    match vak.export_bit:
        case ExportModes.no_export.value:
            return
        case ExportModes.table.value:
            if not vak.h5_toetsen:
                return
            # otherwise, do the old behaviour after the pattern match
        case ExportModes.remark_completed_earlier.value:
            add_header(document, vak, year, leerjaar)
            paragraph = document.add_paragraph(
                "De leerling heeft dit vak in een eerder stadium afgesloten op Vwo "
                "niveau"
            )
            # try to set global font name
            _set_default_font(paragraph)
            document.add_page_break()
            return
        case ExportModes.remark_vwo.value:
            add_header(document, vak, year, leerjaar)
            paragraph = document.add_paragraph(
                "De leerling rond dit vak af op Vwo niveau. Eventuele openstaande "
                "opdrachten in het Vwo 6 PTA moeten worden afgerond."
            )
            # try to set global font name
            _set_default_font(paragraph)
            document.add_page_break()
            return

    if not vak.h5_toetsen:
        return

    overstappen = {
        overstap.oude_toets_id: overstap for overstap in vak.overstappen_vwo6
    }

    add_header(document, vak, year, leerjaar)

    paragraph = document.add_paragraph(
        "De volgende al gemaakte onderdelen worden meegenomen uit eerdere leerjaren.\n"
        "Daarbij wordt het cijfer overgenomen, of de toets wordt opnieuw gewaardeerd.\n"
        "Toetsen van periode 3 worden meegemaakt met de Havo 5 klas."
    )
    # try to set global font name
    _set_default_font(paragraph)

    PREFIX_MAPPING = {
        Leerjaren.vwo_4: "V4",
        Leerjaren.vwo_5: "V5",
        Leerjaren.vwo_6: "V6",
    }

    def _get_oude_toets(toets) -> str:
        overstap = overstappen.get(toets.id)
        if not overstap:
            return ""
        if not overstap.h5_toets:
            return ""
        prefix = PREFIX_MAPPING[overstap.h5_toets.klas]
        return f"{prefix} {overstap.h5_toets.code}"

    def _get_actie(toets):
        overstap = overstappen.get(toets.id)
        if not overstap:
            return ""
        return overstap.get_actie_display()

    table_data = [
        [
            toets.code,
            _get_oude_toets(toets),
            clean_text(toets.omschrijving),
            toets.domein or "",
            _get_actie(toets),
        ]
        for toets in vak.h5_toetsen
    ]

    header = [
        "Code H5",
        "Oude toets",
        "Omschrijving",
        "Domein",
        "Actie",
    ]
    WIDTHS = {
        0: Cm(2.00),
        1: Cm(2.50),
        2: Cm(10.43),
        3: Cm(3.25),
        4: Cm(3.00),
    }
    create_table(document, header, table_data, index_no_center=2, widths=WIDTHS)

    document.add_page_break()


def add_vak_tl(
    document: Document,
    vak: Vak,
    year: int,
    leerjaar: int,
    toetsweek_periodes: dict[int, list[int]],
) -> None:
    toetsen: list[Toets]
    if not (toetsen := vak.toetsen):
        return

    add_header(document, vak, year, leerjaar)

    type_map = {
        3: "schriftelijk",
        4: "praktisch",
        5: "handelingsdeel",
    }

    weging_map = {
        Leerjaren.tl_3: lambda t: t.weging_ed3,
        Leerjaren.tl_4: lambda t: t.weging_ed4,
    }

    table_data = [
        [
            toets.code,
            clean_text(toets.omschrijving),
            toets.domein or "",
            get_periode_and_week(toets, toetsweek_periodes)[0],
            toets.get_herkansbaar_display(),
            type_map.get(toets.type) or "",  # does not follow regular enum...
            str(toets.tijd) or "",
            str(weging_map[leerjaar](toets)) if leerjaar in weging_map else "",
        ]
        for toets in toetsen
    ]

    header = [
        "Code",
        "Onderwerp/Omschrijving",
        "Eind-\ntermen",
        "Periode",
        "Herkansbaar",
        "Soort werk",
        "Tijd\n(min)",
        "Weging SE",
    ]
    WIDTHS = {
        0: Cm(1.51),
        1: Cm(10.43),
        2: Cm(2.00),
        3: Cm(1.50),
        4: Cm(2.00),
        5: Cm(3.00),
        6: Cm(1.50),
        7: Cm(2.50),
    }

    create_table(document, header, table_data, index_no_center=1, widths=WIDTHS)

    document.add_page_break()


def create_table(
    document, header, table_data, widths: dict[int, Cm], index_no_center: int = 1
):
    table = document.add_table(rows=len(table_data) + 1, cols=len(header))
    _set_default_table_style(table)

    # add table header
    _set_table_header(table, header)

    # add table body
    for row_index, data_row in enumerate(table_data, 1):
        for column_index, content in enumerate(data_row):
            cell = table.cell(row_index, column_index)
            cell.text = content

    _style_table_cells(table, index_no_center=index_no_center)
    _set_column_widths(table, widths)

    return table


def _set_default_font(paragraph) -> None:
    paragraph.style.font.name = "Arial"
    paragraph.style.font.size = Pt(10)


def _set_default_table_style(table) -> None:
    table.style = "TableGrid"
    table.autofit = False

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.7)

        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_table_cells(table, index_no_center=1):
    # style header row
    _set_table_header_style(table.rows[0], index_no_center=index_no_center)

    # style body
    for row in table.rows:
        # style the cells
        for _index, cell in enumerate(row.cells):
            par = cell.paragraphs[0]
            # first column bold
            if _index == 0:
                for run in par.runs:
                    run.font.bold = True
            # center if needed
            if _index != index_no_center:
                for par in cell.paragraphs:
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_table_header(table, header: List[str]) -> None:
    header_cells = table.rows[0].cells
    for index, title in enumerate(header):
        cell = header_cells[index]
        cell.text = title


def _set_table_header_style(header_row, index_no_center=1) -> None:
    for index, cell in enumerate(header_row.cells):
        cell.paragraphs[0].runs[0].bold = True
        if index != index_no_center:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell)


def _set_column_widths(table, widths) -> None:
    for index, column in enumerate(table.columns):
        column.width = widths[index]
        # sigh... dumb format
        for cell in column.cells:
            cell.width = column.width


def _set_cell_bg(cell, color: str = HEADER_BG_COLOR):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)
